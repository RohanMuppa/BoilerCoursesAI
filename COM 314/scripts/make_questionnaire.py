from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', 'scripts'))
from google_drive_api import GoogleDriveAPI

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

sections = doc.sections
for s in sections:
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

# Speaker number top right
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('#9')
run.bold = True
run.font.size = Pt(24)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Audience Analysis Questionnaire')
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COM 314/315')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.space_after = Pt(6)
run = p.add_run('Thank you for taking the time to complete this questionnaire. Your honest responses will help me select and develop topics for upcoming presentations. All responses are anonymous.')
run.font.size = Pt(10)
run.italic = True

def add_topic_header(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.underline = True

def add_question(doc, num, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    run = p.add_run(f'{num}. {text}')
    run.bold = True
    run.font.size = Pt(11)

def add_option(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(f'\u25cb  {text}')
    run.font.size = Pt(11)

def add_scale(doc, left_label, right_label):
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(f'{left_label}    1        2        3        4        5    {right_label}')
    run.font.size = Pt(11)

def add_blank_line(doc):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run('_' * 70)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 180)

# TOPIC 1
add_topic_header(doc, 'TOPIC 1: AI, Automation, and Future Careers')

add_question(doc, 1, 'What is your major or intended field of study?')
add_blank_line(doc)

add_question(doc, 2, 'How much do you know about how AI and automation are expected to change the job market for college graduates?')
add_option(doc, 'I know a lot about this topic')
add_option(doc, 'I know a little about this topic')
add_option(doc, 'I have heard of it but do not know much')
add_option(doc, 'I have never heard of this topic')

add_question(doc, 3, 'How concerned are you about AI or automation affecting your own future career?')
add_scale(doc, 'Not concerned', 'Very concerned')

add_question(doc, 4, 'Have you ever changed or reconsidered your career plans because of advancements in AI or automation?')
add_option(doc, 'Yes')
add_option(doc, 'No')
add_option(doc, 'Unsure')

add_question(doc, 5, 'What questions or concerns do you have about AI\'s impact on your future career?')
add_blank_line(doc)
add_blank_line(doc)

# TOPIC 2
add_topic_header(doc, 'TOPIC 2: AGI and the Future of Ownership, Work, and Autonomy')

add_question(doc, 6, 'Have you heard of the concept of Artificial General Intelligence (AGI)?')
add_option(doc, 'Yes, I understand what it means')
add_option(doc, 'I have heard of it but I am not sure what it means')
add_option(doc, 'No, I have never heard of it')

add_question(doc, 7, 'Are you familiar with the phrase "you will own nothing and be happy"?')
add_option(doc, 'Yes')
add_option(doc, 'No')

add_question(doc, 8, 'How interested would you be in learning about how AGI could reshape ownership, work, and personal freedom in the future?')
add_scale(doc, 'Not interested', 'Very interested')

add_question(doc, 9, 'How much do you agree with the following statement: "Technology companies have too much control over people\'s daily lives."')
add_scale(doc, 'Strongly disagree', 'Strongly agree')

add_question(doc, 10, 'What comes to mind when you think about AGI and its potential impact on society?')
add_blank_line(doc)
add_blank_line(doc)

# TOPIC 3
add_topic_header(doc, 'TOPIC 3: Personal AI Agents and Tools for College Students')

add_question(doc, 11, 'Do you currently use any AI tools (such as ChatGPT, Copilot, Grammarly, or similar) in your daily life or schoolwork?')
add_option(doc, 'Yes, regularly')
add_option(doc, 'Yes, occasionally')
add_option(doc, 'No, but I have tried them')
add_option(doc, 'No, I have never used them')

add_question(doc, 12, 'How comfortable are you with the idea of using AI agents to manage tasks like scheduling, studying, or organizing your personal life?')
add_scale(doc, 'Not comfortable', 'Very comfortable')

add_question(doc, 13, 'Do you think college students should learn to build or customize their own AI tools?')
add_option(doc, 'Yes, it is an important skill for the future')
add_option(doc, 'Maybe, depending on the field of study')
add_option(doc, 'No, it is not necessary')
add_option(doc, 'I am not sure')

add_question(doc, 14, 'How interested would you be in a presentation about how students can use personal AI agents to take more control of their lives?')
add_scale(doc, 'Not interested', 'Very interested')

# TOPIC 4
add_topic_header(doc, 'TOPIC 4: Philosophy of Education and the College Experience')

add_question(doc, 15, 'Have you ever taken a philosophy course or read philosophical works on your own?')
add_option(doc, 'Yes')
add_option(doc, 'No')

add_question(doc, 16, 'How often do you reflect on what you personally want to get out of your college experience beyond earning a degree?')
add_option(doc, 'Often')
add_option(doc, 'Sometimes')
add_option(doc, 'Rarely')
add_option(doc, 'Never')

add_question(doc, 17, 'How interested would you be in a presentation exploring philosophical ideas about what education should really be for?')
add_scale(doc, 'Not interested', 'Very interested')

add_question(doc, 18, 'How much do you agree with the following statement: "College is more about personal growth than career preparation."')
add_scale(doc, 'Strongly disagree', 'Strongly agree')

add_question(doc, 19, 'What do you hope to gain from your college experience beyond a degree or job qualification?')
add_blank_line(doc)
add_blank_line(doc)

add_question(doc, 20, 'Of the four topics below, which would you be most interested in hearing a presentation about? (Circle one)')
add_option(doc, 'How AI and automation are changing future careers')
add_option(doc, 'How AGI could reshape ownership, work, and personal freedom')
add_option(doc, 'How college students can use personal AI agents in daily life')
add_option(doc, 'Philosophical perspectives on education and personal growth')

# Thank you
p = doc.add_paragraph()
p.space_before = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Thank you for your responses.')
run.italic = True
run.font.size = Pt(11)

outpath = '/Users/rohanmuppa/Documents/school/COM 314/Audience Analysis Questionnaire.docx'
doc.save(outpath)
print(f'Saved to {outpath}')

# Upload to Google Drive
drive = GoogleDriveAPI()
drive.authenticate()

result = drive.upload_file(outpath, mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
print(f'Uploaded to Google Drive: {result}')
